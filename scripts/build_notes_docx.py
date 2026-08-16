#!/usr/bin/env python3
"""Build a compact, mobile-readable Word note from structured UTF-8 JSON."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Microsoft YaHei"
NAVY = RGBColor(32, 55, 72)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(96, 104, 112)
LIGHT_FILL = "F4F6F9"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="UTF-8 notes JSON")
    parser.add_argument("--output", required=True, type=Path, help="Output DOCX")
    return parser.parse_args()


def set_run_font(run, size: float | None = None, color=None, bold=None, italic=None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = FONT
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = NAVY
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def shade_paragraph(paragraph, fill: str = LIGHT_FILL) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def clean_source(value: str) -> str:
    value = value.strip()
    value = re.sub(r"([?&])(token|access_token|auth|signature|sig|key|pwd)=[^&#\s]+", "", value, flags=re.I)
    return value.rstrip("?&")


def add_paragraphs(doc: Document, text: str) -> None:
    chunks = [part.strip() for part in re.split(r"\n\s*\n", text.strip()) if part.strip()]
    for chunk in chunks:
        doc.add_paragraph(chunk)


def validate(data: dict) -> None:
    for field in ("title", "one_sentence", "summary", "sections", "time_ranges"):
        if not data.get(field):
            raise ValueError(f"Missing required field: {field}")
    if not isinstance(data["summary"], list) or not all(isinstance(x, str) and x.strip() for x in data["summary"]):
        raise ValueError("summary must be a non-empty array of strings")
    if not isinstance(data["sections"], list):
        raise ValueError("sections must be an array")
    for index, section in enumerate(data["sections"], start=1):
        if not section.get("heading") or not (section.get("paragraphs") or section.get("bullets")):
            raise ValueError(f"section {index} needs heading and content")
    if not isinstance(data["time_ranges"], list) or len(data["time_ranges"]) != 3:
        raise ValueError("time_ranges must contain exactly three entries")
    for item in data["time_ranges"]:
        if not item.get("range") or not item.get("summary"):
            raise ValueError("each time range needs range and summary")
    if str(data.get("language", "")).lower().startswith("en") and not str(data.get("english_transcript", "")).strip():
        raise ValueError("English videos require english_transcript")


def build(data: dict, output: Path) -> None:
    validate(data)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    doc.core_properties.title = str(data["title"])
    doc.core_properties.subject = "视频转 Word 笔记"

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("视频学习笔记"), size=8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("由视频内容整理 · 请结合原视频校核"), size=8.5, color=GRAY)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(str(data["title"]))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details = []
    if data.get("duration"):
        details.append(f"时长：{data['duration']}")
    if data.get("language"):
        details.append(f"语言：{data['language']}")
    set_run_font(subtitle.add_run("  |  ".join(details) or "结构化视频笔记"), size=9.5, color=GRAY)
    if data.get("source"):
        source_p = doc.add_paragraph()
        source_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(source_p.add_run(f"来源：{clean_source(str(data['source']))}"), size=8.5, color=GRAY)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(12)
    lead.paragraph_format.space_after = Pt(12)
    shade_paragraph(lead)
    set_run_font(lead.add_run("一句话结论\n"), size=10, color=BLUE, bold=True)
    set_run_font(lead.add_run(str(data["one_sentence"])), size=11.5, color=NAVY, bold=True)

    doc.add_heading("核心摘要", level=1)
    for item in data["summary"]:
        doc.add_paragraph(item.strip(), style="List Bullet")

    doc.add_heading("完整笔记", level=1)
    for item in data["sections"]:
        doc.add_heading(str(item["heading"]), level=2)
        for paragraph in item.get("paragraphs", []):
            if str(paragraph).strip():
                doc.add_paragraph(str(paragraph).strip())
        for bullet in item.get("bullets", []):
            if str(bullet).strip():
                doc.add_paragraph(str(bullet).strip(), style="List Bullet")

    if data.get("checklist"):
        doc.add_heading("行动检查清单", level=1)
        for item in data["checklist"]:
            doc.add_paragraph(str(item).strip(), style="List Number")

    if str(data.get("translated_full_text", "")).strip():
        doc.add_heading("中文全文翻译", level=1)
        add_paragraphs(doc, str(data["translated_full_text"]))

    doc.add_heading("三段时间概览", level=1)
    for item in data["time_ranges"]:
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"{item['range']}  "), size=11, color=DARK_BLUE, bold=True)
        set_run_font(p.add_run(str(item["summary"])), size=11)

    if str(data.get("english_transcript", "")).strip():
        doc.add_page_break()
        doc.add_heading("英文全文", level=1)
        note = doc.add_paragraph()
        set_run_font(note.add_run("以下为清理后的英文原文，保留原意并修正明显断句。"), size=9.5, color=GRAY, italic=True)
        add_paragraphs(doc, str(data["english_transcript"]))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    args = parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8-sig"))
    build(data, args.output.expanduser().resolve())
    print(args.output.expanduser().resolve())
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except (ValueError, json.JSONDecodeError) as exc:
        raise SystemExit(f"Invalid notes JSON: {exc}")
